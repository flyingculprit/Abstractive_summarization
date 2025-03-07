from flask import Flask, render_template, request, send_file
import os
from docx import Document
import google.generativeai as genai

# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = "word_uploads"

# Ensure upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Configure Gemini API
genai.configure(api_key="your api")

def extract_text_from_docx(docx_path):
    """Extract text from a Word document (.docx)."""
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def summarize_text(text):
    """Generate summary using Gemini API."""
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content(f"Summarize the following text:\n{text}")
    return response.text.strip() if response else "Error in summarization."

def save_summary_as_docx(summary_text, output_path):
    """Save the summarized text as a Word document."""
    doc = Document()
    doc.add_paragraph(summary_text)
    doc.save(output_path)

@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        file = request.files["file"]

        if file and file.filename.endswith(".docx"):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)

            # Extract text from uploaded Word document
            original_text = extract_text_from_docx(file_path)
            original_word_count = len(original_text.split())

            # Generate summary
            summary_text = summarize_text(original_text)
            summarized_word_count = len(summary_text.split())
            summary_ratio = round((summarized_word_count / original_word_count) * 100, 2)

            # Save summary as a Word document
            summary_filename = "summary.docx"
            summary_path = os.path.join(app.config['UPLOAD_FOLDER'], summary_filename)
            save_summary_as_docx(summary_text, summary_path)

            return render_template("word_result.html", 
                                   original_text=original_text, 
                                   original_word_count=original_word_count,
                                   summary_text=summary_text, 
                                   summarized_word_count=summarized_word_count, 
                                   summary_ratio=summary_ratio, 
                                   summary_file=summary_filename)

    return render_template("word_index.html")

@app.route("/download/<filename>")
def download_file(filename):
    """Allow downloading the summarized Word document."""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
